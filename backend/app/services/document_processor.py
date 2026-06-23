import hashlib
import logging
import re
import pytesseract

pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
)

from app.config.settings import get_settings

logger = logging.getLogger(__name__)

settings = get_settings()


@dataclass
class ExtractedChunk:
    text: str
    metadata: dict[str, Any]


@dataclass
class ExtractionResult:
    text: str
    chunks: list[ExtractedChunk]
    page_count: int
    table_count: int


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for block in iter(
            lambda: file.read(1024 * 1024),
            b"",
        ):
            digest.update(block)
    return digest.hexdigest()


LIGATURE_REPAIRS = {
    r"ﬁ": "fi",
    r"ﬂ": "fl",
    r"ﬀ": "ff",
    r"ﬃ": "ffi",
    r"ﬄ": "ffl",
}

def clean_text(text: str) -> str:
    if not text:
        return ""

    # Remove markup tags safely
    text = re.sub(r"<[^>]+>", " ", text)
    text = text.replace("\x00", " ")
    
    # Remove markdown line artifacts
    text = re.sub(r"\*{3,}", " ", text)
    text = re.sub(r"_{3,}", " ", text)
    
    # Fix common ligatures and OCR broken words
    for pattern, repaired in LIGATURE_REPAIRS.items():
        text = re.sub(pattern, repaired, text)

    # Fix OCR hyphen line breaks safely
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    # FIXED: Only clean hyphens if they are trailing on line boundaries 
    # to protect standard compound technical components and identifiers
    text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)

    # Fix underscore-wrapped words (common OCR artifact)
    text = re.sub(r'_(\w+)_', r'\1', text)
    text = re.sub(r'\*(\w+)\*', r'\1', text)

    # Normalize spaces safely without destroying technical layouts
    text = re.sub(r"[ \t]+", " ", text)

    # Filter out long decorative lines or divider artifacts common in NRSC reports
    text = re.sub(r"[-_.*=]{4,}", " ", text)

    # Normalize repeated characters safely
    text = re.sub(r"([^`*#\s])\1{6,}", r"\1", text)

    # Ensure clean trailing white-space lines are fixed
    text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)

    # Standardize empty paragraph spacing block margins
    text = re.sub(r"\n{3,}", "\n\n", text)

    # FIXED: Removed the destructive CamelCase character splitter regex completely 
    # to protect framework descriptors and variable definitions unharmed.

    return text.strip()


def is_meaningful_text(text: str) -> bool:
    if not text:
        return False

    text = text.strip()

    # Minimum threshold to catch tiny sentences with real mathematical significance
    if len(text) < 15:
        return False

    # Skip validation for structural formatting elements
    if "[TABLE]" in text or "```" in text:
        return True

    # Loosened special character ratio to protect expressions and markdown grids
    special_ratio = (
        len(re.findall(r"[^a-zA-Z0-9\s.,:;()\-+\*/%=<>|]", text))
        / max(len(text), 1)
    )

    if special_ratio > 0.35:
        return False

    # Ensure chunks contain realistic vocabulary thresholds
    words = re.findall(r"\b[a-zA-Z0-9_]{2,}\b", text)
    if len(words) < 4:
        return False

    garbage_patterns = [
        r"[^\s]{55,}",  # Extended threshold limit safely to keep long hashes intact
    ]

    for pattern in garbage_patterns:
        if re.search(pattern, text):
            return False

    return True


def extract_docx(path: Path) -> tuple[list[tuple[int, str]], int]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    parts = []
    table_count = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        table_count += 1
        rows = []
        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]
            rows.append(" | ".join(cells))

        parts.append("\n\n[TABLE]\n" + "\n".join(rows))

    return [(1, "\n\n".join(parts))], table_count


def extract_images_from_pdf_page(page) -> str:
    extracted_text = []
    try:
        images = page.images
        if not images:
            return ""

        for _ in images:
            try:
                image = page.to_image(resolution=250).original
                text = pytesseract.image_to_string(image)
                text = clean_text(text)

                if is_meaningful_text(text):
                    extracted_text.append(text)
            except Exception:
                continue
    except Exception:
        return ""

    return "\n".join(extracted_text)


def extract_pdf_text(path: Path) -> tuple[list[tuple[int, str]], int]:
    pages = []
    table_count = 0

    try:
        import fitz
    except ImportError:
        fitz = None

    if fitz is not None:
        try:
            pdf = fitz.open(path)
            for page_num, page in enumerate(pdf, start=1):
                raw_text = page.get_text("text") or ""
                
                # FIXED: Check if PyMuPDF natively found tabular blocks first
                try:
                    tabs = page.find_tables()
                    if tabs:
                        table_text_segments = []
                        for tab in tabs:
                            table_count += 1
                            table_data = tab.extract()
                            table_lines = [" | ".join([str(cell or "").strip().replace("\n", " ") for cell in r]) for r in table_data]
                            table_text_segments.append("\n".join(table_lines))
                        
                        # Add structural tags smoothly
                        raw_text += "\n\n[TABLE]\n" + "\n\n".join(table_text_segments)
                except Exception:
                    pass

                text = clean_text(raw_text)

                # Fallback to OCR if page content is heavily embedded imagery
                if len(text.strip()) < 40 or not is_meaningful_text(text):
                    try:
                        pix = page.get_pixmap(dpi=200)
                        import PIL.Image
                        import io

                        image = PIL.Image.open(io.BytesIO(pix.tobytes("png")))
                        ocr_text = pytesseract.image_to_string(image)
                        ocr_text = clean_text(ocr_text)

                        if is_meaningful_text(ocr_text):
                            text = ocr_text
                    except Exception:
                        pass

                pages.append((page_num, text))
            pdf.close()
        except Exception:
            logger.exception("PyMuPDF extraction failed for %s; trying fallback engines.", path)

    if pages:
        return pages, table_count

    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = clean_text(page.extract_text() or "")
                
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []

                table_lines = []
                for table in tables:
                    table_count += 1
                    for row in table:
                        cells = [clean_text(cell or "") for cell in row]
                        table_lines.append(" | ".join(cells))

                if table_lines:
                    text = (text + "\n\n[TABLE]\n" + "\n".join(table_lines)).strip()

                pages.append((page_num, text))
    except Exception:
        logger.exception("pdfplumber extraction failed for %s; trying PyPDF2 fallback.", path)

    if pages:
        return pages, table_count

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        for page_num, page in enumerate(reader.pages, start=1):
            pages.append((page_num, clean_text(page.extract_text() or "")))
    except Exception:
        logger.exception("PyPDF2 extraction failed completely for %s", path)

    return pages, table_count


def extract_pdf_ocr(path: Path, existing_pages: list[tuple[int, str]]) -> list[tuple[int, str]]:
    try:
        from pdf2image import convert_from_path
    except Exception:
        logger.warning("OCR libraries unavailable for %s", path)
        return existing_pages

    page_map = {page: text for page, text in existing_pages}

    try:
        images = convert_from_path(
            str(path),
            dpi=150,
            thread_count=2,
        )

        for idx, image in enumerate(images, start=1):
            existing = page_map.get(idx, "") or ""
            if is_meaningful_text(existing) and len(existing.strip()) > 50:
                continue

            ocr_text = pytesseract.image_to_string(image) or ""
            ocr_text = clean_text(ocr_text)

            if not is_meaningful_text(ocr_text):
                continue

            if len(ocr_text) > len(existing) * 1.2:
                page_map[idx] = ocr_text
    except Exception:
        logger.exception("OCR pipeline fallback failed for %s", path)

    return sorted(page_map.items(), key=lambda item: item[0])


def extract_txt(path: Path) -> tuple[list[tuple[int, str]], int]:
    return [(1, path.read_text(encoding="utf-8", errors="ignore"))], 0


def deduplicate_chunks(chunks: list[ExtractedChunk]) -> list[ExtractedChunk]:
    unique = []
    seen = set()

    for chunk in chunks:
        # Standardize space configurations to hash contents accurately
        normalized = re.sub(r"\s+", " ", chunk.text).strip().lower()
        key = hashlib.md5(normalized.encode()).hexdigest()
        if key in seen:
            continue
        seen.add(key)
        unique.append(chunk)

    return unique


def classify_chunk_kind(text: str) -> str:
    if "[TABLE]" in text:
        return "table"
    if len(text.split()) < 50:
        return "short"
    return "general"


def extract_document(path: Path, filename: str, document_type: str) -> ExtractionResult:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        pages, table_count = extract_pdf_text(path)
        pages = extract_pdf_ocr(path, pages)
    elif suffix == ".docx":
        pages, table_count = extract_docx(path)
    elif suffix in {".txt", ".md", ".csv"}:
        pages, table_count = extract_txt(path)
    else:
        raise ValueError("Unsupported document type")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=[
            "\n# ", "\n## ", "\n### ",
            "\n#### ", "\n##### ",
            "\n[TABLE]", "\n[IMAGE_TEXT]", "\n```",
            "\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""
        ],
        keep_separator=True,
    )

    chunks = []
    all_text = []

    for page, raw in pages:
        cleaned_raw = clean_text(raw)
        if not cleaned_raw:
            continue

        all_text.append(f"[Page {page}]\n{cleaned_raw}")

        split_chunks = splitter.split_text(cleaned_raw)
        
        for idx, chunk in enumerate(split_chunks, start=1):
            chunk_content = chunk.strip()
            if len(chunk_content) < 40:
                continue

            # Ensure valid validation sequences run cleanly
            if (
                not is_meaningful_text(chunk_content)
                and "[TABLE]" not in chunk_content
                and "```" not in chunk_content
            ):
                continue

            # Detect section headings for rich metadata
            heading_match = re.search(r'^(#{1,5}\s+.*)$', chunk_content, re.MULTILINE)
            section_heading = heading_match.group(1).strip() if heading_match else ""

            chunks.append(
                ExtractedChunk(
                    text=chunk_content,
                    metadata={
                        "file": filename,
                        "filename": filename,
                        "page": page,
                        "chunk_id": f"{path.stem}-p{page}-c{idx}",
                        "document_type": document_type,
                        "chunk_kind": classify_chunk_kind(chunk_content),
                        "section_heading": section_heading,
                        "chunk_index": idx,
                        "total_chunks": len(split_chunks),
                    },
                )
            )

    chunks = deduplicate_chunks(chunks)

    if not chunks:
        raise ValueError("No searchable text extracted")

    logger.info(f"Extracted {len(chunks)} chunks from {filename}")

    return ExtractionResult(
        text="\n\n".join(all_text),
        chunks=chunks,
        page_count=max(len(pages), 1),
        table_count=table_count,
    )